from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUT = DOCS / "Almodiel_Trucking_Services_System_Documentation.docx"
ERD_IMAGE = DOCS / "ERD-Access-Style-wide.png"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    run.bold = bold


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2.0
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Inches(0.5)
    normal.paragraph_format.space_after = Pt(0)

    for style_name in ("Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(12)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = 2.0
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)

    for sec in doc.sections:
        add_page_number(sec.footer.paragraphs[0])


def add_centered(doc, text, bold=False, size=12, after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.first_line_indent = None
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {min(level, 4)}")
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.line_spacing = 2.0
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(text).bold = True
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text)
    return p


def add_bullet(doc, term, definition):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(term + ": ")
    r.bold = True
    p.add_run(definition)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        set_cell_text(hdr.cells[i], h, bold=True)
        set_cell_shading(hdr.cells[i], "D9EAF7")
        hdr.cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            set_cell_width(hdr.cells[i], widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                set_cell_width(cells[i], widths[i])
    doc.add_paragraph()
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.first_line_indent = None
    r = p.add_run(text)
    r.bold = True


def field_rows(fields):
    return [[name, typ, key, fmt, desc] for name, typ, key, fmt, desc in fields]


TABLES = {
    "booking": {
        "description": "The booking entity stores the main delivery request record. It connects customers, pickup and destination locations, pricing, trip grouping, and delivery status.",
        "source": "Booking Registration Form",
        "processes": "Create Booking, Update Trip, Start Delivery, Complete Delivery, Generate Sales",
        "triggered": "Admin, Driver, Customer",
        "fields": [
            ("bookingID", "Number", "PK", "Auto-number", "Unique booking identifier."),
            ("customerID", "Number", "FK/logical", "N/A", "References the customer who owns the booking."),
            ("storeName", "Short Text(150)", "", "N/A", "Stores the customer or store display name for the booking."),
            ("pickupLocationID", "Number", "FK/logical", "N/A", "References the pickup location."),
            ("destinationLocationID", "Number", "FK/logical", "N/A", "References the destination location."),
            ("tripID", "Number", "Logical key", "N/A", "Groups one or more booking records into one trip."),
            ("pickupDateTime", "Date/Time", "", "MM/DD/YYYY HH:MM", "Scheduled pickup date and time."),
            ("price", "Currency/Number", "", "0.00", "Booking amount charged to the customer."),
            ("createdBy", "Number", "FK/logical", "N/A", "Employee or admin who created the booking."),
            ("dateCreated", "Date/Time", "", "MM/DD/YYYY HH:MM", "Date and time the record was created."),
            ("status", "Short Text(20)", "", "pending | in-transit | stopover | completed", "Current delivery status."),
        ],
    },
    "cargo": {
        "description": "The cargo entity stores one or more cargo items per booking. It allows the system to document cargo type, quantity, condition, and handling needs.",
        "source": "Booking Registration Form",
        "processes": "Add Cargo, Review Booking, View Trip Details",
        "triggered": "Admin, Customer",
        "fields": [
            ("cargoID", "Number", "PK", "Auto-number", "Unique cargo identifier."),
            ("bookingID", "Number", "FK/logical", "N/A", "Booking where the cargo belongs."),
            ("cargoType", "Short Text(100)", "", "N/A", "Type or name of cargo."),
            ("quantity", "Number", "", "Whole number", "Quantity of the cargo item."),
            ("condition", "Short Text(100)", "", "N/A", "Condition of the cargo."),
            ("description", "Long Text", "", "N/A", "Detailed cargo description."),
            ("specialHandling", "Long Text", "", "N/A", "Special instructions for handling the cargo."),
        ],
    },
    "customer": {
        "description": "The customer entity stores individual and company customer accounts. It supports customer identification, contact details, registration, and customer-owned bookings.",
        "source": "Customer Registration and Signup Forms",
        "processes": "Register Customer, Manage Company, Create Booking",
        "triggered": "Admin, Customer",
        "fields": [
            ("id", "Number", "PK", "Auto-number", "Unique customer identifier."),
            ("customerType", "Short Text", "", "individual | company", "Identifies whether the customer is individual or company."),
            ("customerFName", "Short Text(100)", "", "N/A", "Customer first name or company name field."),
            ("customerLName", "Short Text(50)", "", "N/A", "Customer last name."),
            ("customerMI", "Short Text(1)", "", "A-Z", "Customer middle initial."),
            ("contactPerson", "Short Text(100)", "", "N/A", "Contact person for company customers."),
            ("email", "Short Text(100)", "", "name@email.com", "Customer email address."),
            ("phoneNumber", "Short Text(11)", "", "09xxxxxxxxx", "Customer phone number."),
            ("province", "Short Text(50)", "", "N/A", "Customer province."),
            ("warehouseLatitude", "Number", "", "Decimal", "Warehouse latitude for company pickup."),
            ("warehouseLongitude", "Number", "", "Decimal", "Warehouse longitude for company pickup."),
            ("companyDocument", "Short Text(255)", "", "File path", "Uploaded company document path."),
            ("password", "Short Text(255)", "", "Encrypted/plain depending record", "Customer login password."),
            ("dateRegistered", "Date/Time", "", "MM/DD/YYYY", "Customer registration date."),
            ("status", "Short Text", "", "active | inactive", "Customer account status."),
            ("locationID", "Number", "FK", "N/A", "Linked warehouse or customer location."),
        ],
    },
    "deliverycharge": {
        "description": "The deliverycharge entity stores hauling and other charges connected to a booking and trip. It supports billing and operational charge tracking.",
        "source": "Booking Registration and Delivery Charge Forms",
        "processes": "Record Hauling Charge, Review Billing, Generate Reports",
        "triggered": "Admin",
        "fields": [
            ("deliveryChargeID", "Number", "PK", "Auto-number", "Unique delivery charge identifier."),
            ("bookingID", "Number", "FK/logical", "N/A", "Booking charged."),
            ("tripID", "Number", "Logical key", "N/A", "Trip group charged."),
            ("chargeType", "Short Text", "", "hauling | others", "Type of delivery charge."),
            ("amount", "Currency/Number", "", "0.00", "Charge amount."),
            ("notes", "Long Text", "", "N/A", "Charge notes."),
            ("createdBy", "Number", "FK/logical", "N/A", "User who recorded the charge."),
            ("dateCreated", "Date/Time", "", "MM/DD/YYYY HH:MM", "Date created."),
        ],
    },
    "employee": {
        "description": "The employee entity stores staff accounts for admins, drivers, and assistants. It supports authentication, crew assignment, salary processing, and incident accountability.",
        "source": "Employee Registration Form",
        "processes": "Register Employee, Assign Crew, Login, Generate Salary",
        "triggered": "Admin",
        "fields": [
            ("id", "Number", "PK", "Auto-number", "Unique employee identifier."),
            ("empFName", "Short Text(50)", "", "N/A", "Employee first name."),
            ("empLName", "Short Text(50)", "", "N/A", "Employee last name."),
            ("empMI", "Short Text(1)", "", "A-Z", "Middle initial."),
            ("empSuffix", "Short Text(10)", "", "N/A", "Name suffix."),
            ("empBirthDate", "Date/Time", "", "MM/DD/YYYY", "Birth date."),
            ("empPhoneNumber", "Short Text(20)", "", "09xxxxxxxxx", "Phone number used for login."),
            ("empEmail", "Short Text(100)", "", "name@email.com", "Email address."),
            ("empType", "Short Text", "", "admin | driver | assistant", "Employee role type."),
            ("empStatus", "Short Text", "", "active | inactive", "Employee status."),
            ("dateCreated", "Date/Time", "", "MM/DD/YYYY HH:MM", "Date registered."),
            ("empPassword", "Short Text(255)", "", "N/A", "Employee login password."),
            ("licenseNumber", "Short Text(50)", "", "N/A", "Driver license number."),
            ("licenseExpire", "Short Text(50)", "", "MM/DD/YYYY", "License expiration date."),
            ("licenseImage", "Short Text(255)", "", "File path", "Uploaded license image."),
        ],
    },
    "expenses": {
        "description": "The expenses entity stores operational expenses such as fuel, maintenance, salary, tolls, parking, repairs, office costs, and other costs.",
        "source": "Expense Record Form and System-generated Salary/Fuel Records",
        "processes": "Record Expense, Review Reports, Compute Net Sales",
        "triggered": "Admin, System",
        "fields": [
            ("expenseID", "Number", "PK", "Auto-number", "Unique expense identifier."),
            ("expenseDate", "Date/Time", "", "MM/DD/YYYY", "Date of expense."),
            ("category", "Short Text", "", "fuel | truck_maintenance | employee_salary | truck_document | toll | parking | repair | office | other", "Expense category."),
            ("amount", "Currency/Number", "", "0.00", "Expense amount."),
            ("description", "Long Text", "", "N/A", "Expense description."),
            ("truckID", "Number", "FK", "N/A", "Related truck."),
            ("empID", "Number", "FK", "N/A", "Related employee."),
            ("tripID", "Number", "Logical key", "N/A", "Related trip."),
            ("bookingID", "Number", "FK/logical", "N/A", "Related booking."),
            ("referenceNo", "Short Text(100)", "", "N/A", "Receipt or reference number."),
            ("receiptImage", "Short Text(255)", "", "File path", "Uploaded receipt image."),
            ("status", "Short Text", "", "pending | approved | paid | cancelled", "Expense status."),
            ("createdBy", "Number", "FK", "N/A", "Admin/user who created the expense."),
            ("dateCreated", "Date/Time", "", "MM/DD/YYYY HH:MM", "Date created."),
        ],
    },
    "incidentreport": {
        "description": "The incidentreport entity stores driver-submitted incidents and admin review details. It supports safety monitoring, accountability, and resolution tracking.",
        "source": "Incident Report Form",
        "processes": "Submit Incident, Review Incident, Resolve Incident",
        "triggered": "Driver, Admin",
        "fields": [
            ("incidentID", "Number", "PK", "Auto-number", "Unique incident identifier."),
            ("tripID", "Number", "Logical key", "N/A", "Related trip."),
            ("bookingID", "Number", "FK/logical", "N/A", "Specific booking involved, if any."),
            ("driverID", "Number", "FK/logical", "N/A", "Driver who submitted the report."),
            ("incidentType", "Short Text", "", "accident | vehicle_breakdown | cargo_damage | delay | route_issue | customer_issue | other", "Type of incident."),
            ("severity", "Short Text", "", "low | medium | high | critical", "Incident severity."),
            ("incidentDateTime", "Date/Time", "", "MM/DD/YYYY HH:MM", "Date and time of incident."),
            ("locationText", "Short Text(255)", "", "N/A", "Incident location."),
            ("description", "Long Text", "", "N/A", "Incident details."),
            ("actionTaken", "Long Text", "", "N/A", "Action taken by driver/admin."),
            ("status", "Short Text", "", "open | reviewing | resolved | dismissed", "Review status."),
            ("adminNotes", "Long Text", "", "N/A", "Admin notes."),
            ("reviewedBy", "Number", "FK/logical", "N/A", "Admin who reviewed the report."),
            ("dateSubmitted", "Date/Time", "", "MM/DD/YYYY HH:MM", "Submission date."),
            ("dateUpdated", "Date/Time", "", "MM/DD/YYYY HH:MM", "Last update date."),
        ],
    },
    "location": {
        "description": "The location entity stores structured address and coordinate data. It supports mapping, customer warehouse locations, pickup pins, and destination pins.",
        "source": "Customer Registration and Booking Location Forms",
        "processes": "Save Location, Pin Map, Display Route",
        "triggered": "Admin, Customer",
        "fields": [
            ("locationID", "Number", "PK", "Auto-number", "Unique location identifier."),
            ("province", "Short Text(100)", "", "N/A", "Province."),
            ("city", "Short Text(100)", "", "N/A", "City."),
            ("barangay", "Short Text(100)", "", "N/A", "Barangay."),
            ("street", "Short Text(100)", "", "N/A", "Street."),
            ("description", "Long Text", "", "N/A", "Full address or notes."),
            ("latitude", "Number", "", "Decimal", "Map latitude."),
            ("longitude", "Number", "", "Decimal", "Map longitude."),
        ],
    },
    "sales": {
        "description": "The sales entity stores billing and sales records generated from completed bookings. It supports payment monitoring, net income computation, and sales reports.",
        "source": "Completed Booking and Sales Page",
        "processes": "Generate Sales, Update Payment, Filter Sales, Review Billing",
        "triggered": "Admin, System",
        "fields": [
            ("salesID", "Number", "PK", "Auto-number", "Unique sales identifier."),
            ("bookingID", "Number", "FK", "N/A", "Related booking."),
            ("tripID", "Number", "Logical key", "N/A", "Related trip."),
            ("customerID", "Number", "FK", "N/A", "Related customer."),
            ("grossAmount", "Currency/Number", "", "0.00", "Total sales amount."),
            ("expenseAmount", "Currency/Number", "", "0.00", "Expense amount deducted."),
            ("netAmount", "Currency/Number", "", "0.00", "Net sales amount."),
            ("paidAmount", "Currency/Number", "", "0.00", "Amount paid."),
            ("balanceAmount", "Currency/Number", "", "0.00", "Remaining balance."),
            ("customerType", "Short Text(20)", "", "individual | company", "Customer type."),
            ("paymentStatus", "Short Text(20)", "", "unpaid | partial | paid", "Payment status."),
            ("salesStatus", "Short Text(20)", "", "recorded", "Sales record status."),
            ("dateGenerated", "Date/Time", "", "MM/DD/YYYY HH:MM", "Date generated."),
            ("datePaid", "Date/Time", "", "MM/DD/YYYY HH:MM", "Date paid."),
            ("remarks", "Long Text", "", "N/A", "Sales remarks."),
        ],
    },
    "staffsalary": {
        "description": "The staffsalary entity stores trip-based salary and allowance records for drivers and assistants. It supports payroll monitoring and salary reports.",
        "source": "Booking Crew Salary Fields and Salary Record Form",
        "processes": "Generate Crew Salary, Mark Salary Paid, Review Salary Report",
        "triggered": "Admin, System",
        "fields": [
            ("salaryID", "Number", "PK", "Auto-number", "Unique salary identifier."),
            ("empID", "Number", "FK", "N/A", "Employee being paid."),
            ("tripID", "Number", "Logical key", "N/A", "Related trip."),
            ("creditedBookingID", "Number", "FK", "N/A", "Booking credited for salary."),
            ("creditedDistanceKm", "Number", "", "0.00", "Distance credited."),
            ("tripRole", "Short Text(50)", "", "driver | assistant", "Trip role."),
            ("payPeriodStart", "Date/Time", "", "MM/DD/YYYY", "Pay period start."),
            ("payPeriodEnd", "Date/Time", "", "MM/DD/YYYY", "Pay period end."),
            ("payType", "Short Text", "", "daily | weekly | semi-monthly | monthly | trip | allowance | bonus | adjustment", "Type of pay."),
            ("baseRate", "Currency/Number", "", "0.00", "Base salary or rate."),
            ("grossPay", "Currency/Number", "", "0.00", "Gross pay."),
            ("deductions", "Currency/Number", "", "0.00", "Deductions."),
            ("netPay", "Currency/Number", "", "0.00", "Net pay."),
            ("datePaid", "Date/Time", "", "MM/DD/YYYY HH:MM", "Paid date."),
            ("status", "Short Text", "", "pending | paid | cancelled", "Salary status."),
            ("remarks", "Long Text", "", "N/A", "Salary remarks."),
            ("createdBy", "Number", "FK", "N/A", "User who created salary record."),
            ("dateCreated", "Date/Time", "", "MM/DD/YYYY HH:MM", "Date created."),
        ],
    },
    "tariff": {
        "description": "The tariff entity stores company/customer route pricing rules. It supports tariff matching during booking and calculation of price based on route, truck type, and fuel ranges.",
        "source": "Tariff Management Form",
        "processes": "Create Tariff, Update Tariff, Match Booking Price",
        "triggered": "Admin",
        "fields": [
            ("tariffID", "Number", "PK", "Auto-number", "Unique tariff identifier."),
            ("customerID", "Number", "FK", "N/A", "Customer/company owning the rate."),
            ("branch", "Short Text(100)", "", "N/A", "Branch name."),
            ("origin", "Short Text(100)", "", "N/A", "Route origin."),
            ("destination", "Short Text(255)", "", "N/A", "Route destination."),
            ("distanceKm", "Number", "", "0.00", "Route distance."),
            ("truckType", "Short Text(50)", "", "N/A", "Truck type covered by the rate."),
            ("baseRate", "Currency/Number", "", "0.00", "Base tariff rate."),
            ("hasFuelSubsidy", "Yes/No", "", "0 | 1", "Indicates if fuel subsidy applies."),
            ("fuelRangeStart", "Number", "", "0.00", "Fuel price range start."),
            ("fuelRangeEnd", "Number", "", "0.00", "Fuel price range end."),
            ("fuelSubsidy", "Currency/Number", "", "0.00", "Fuel subsidy amount."),
            ("status", "Short Text", "", "active | inactive", "Tariff status."),
            ("dateCreated", "Date/Time", "", "MM/DD/YYYY HH:MM", "Date created."),
        ],
    },
    "tripemployee": {
        "description": "The tripemployee entity stores driver and assistant assignments for a trip. It connects employees and trucks to the logical trip group.",
        "source": "Booking Registration and Trip Modification Forms",
        "processes": "Assign Crew, Update Crew, Validate Driver Trip Access",
        "triggered": "Admin, System",
        "fields": [
            ("tripEmployeeID", "Number", "PK", "Auto-number", "Unique trip crew identifier."),
            ("tripID", "Number", "Logical key", "N/A", "Related trip group."),
            ("truckID", "Number", "FK/logical", "N/A", "Assigned truck."),
            ("empID", "Number", "FK/logical", "N/A", "Assigned employee."),
            ("role", "Short Text(50)", "", "driver | assistant", "Crew role."),
            ("dateCreated", "Date/Time", "", "MM/DD/YYYY HH:MM", "Assignment date."),
        ],
    },
    "truck": {
        "description": "The truck entity stores truck profile data, current fuel, current mileage, documents, capacity, and status.",
        "source": "Truck Registration Form",
        "processes": "Register Truck, Assign Truck, Update Mileage, View Truck Details",
        "triggered": "Admin, System",
        "fields": [
            ("id", "Number", "PK", "Auto-number", "Unique truck identifier."),
            ("plateNumber", "Short Text(20)", "", "N/A", "Truck plate number."),
            ("type", "Short Text(20)", "", "N/A", "Truck type."),
            ("capacity", "Number", "", "0.00", "Truck capacity."),
            ("fuel", "Number", "", "0", "Current fuel amount."),
            ("mileage", "Number", "", "0", "Current mileage."),
            ("brand", "Short Text(20)", "", "N/A", "Truck brand."),
            ("corDocument", "Short Text(255)", "", "File path", "Certificate of registration file."),
            ("otherDocument", "Short Text(255)", "", "File path", "Other truck document."),
            ("status", "Short Text(20)", "", "active | inactive | in use", "Truck status."),
        ],
    },
    "truckemployee": {
        "description": "The truckemployee entity stores default truck crew assignments. It identifies the regular driver and assistants assigned to a truck.",
        "source": "Truck Registration Form",
        "processes": "Assign Default Crew, View Truck Details",
        "triggered": "Admin",
        "fields": [
            ("truckEmployeeID", "Number", "PK", "Auto-number", "Unique truck crew identifier."),
            ("truckID", "Number", "FK", "N/A", "Related truck."),
            ("empID", "Number", "FK", "N/A", "Related employee."),
            ("role", "Short Text(50)", "", "driver | assistant", "Crew role."),
            ("dateCreated", "Date/Time", "", "MM/DD/YYYY HH:MM", "Assignment date."),
        ],
    },
    "truckfuellog": {
        "description": "The truckfuellog entity stores fuel entries and odometer readings. It supports truck fuel monitoring and expense tracking.",
        "source": "Truck Details Fuel Log Form",
        "processes": "Add Fuel Log, Review Fuel History, Track Truck Fuel",
        "triggered": "Admin, System",
        "fields": [
            ("truckFuelLogID", "Number", "PK", "Auto-number", "Unique fuel log identifier."),
            ("truckID", "Number", "FK/logical", "N/A", "Related truck."),
            ("logDate", "Date/Time", "", "MM/DD/YYYY HH:MM", "Fuel log date."),
            ("litersAdded", "Number", "", "0.00", "Liters added."),
            ("fuelAfter", "Number", "", "0.00", "Fuel after entry."),
            ("odometer", "Number", "", "0.00", "Odometer reading."),
            ("amount", "Currency/Number", "", "0.00", "Fuel cost."),
            ("station", "Short Text(150)", "", "N/A", "Fuel station."),
            ("referenceNo", "Short Text(100)", "", "N/A", "Receipt/reference number."),
            ("notes", "Long Text", "", "N/A", "Fuel log notes."),
            ("createdBy", "Number", "FK/logical", "N/A", "User who created the log."),
            ("dateCreated", "Date/Time", "", "MM/DD/YYYY HH:MM", "Record creation date."),
        ],
    },
    "trucktripusage": {
        "description": "The trucktripusage entity stores fuel and mileage usage after trip completion. It supports automatic truck mileage and fuel updating.",
        "source": "Completed Trip Status Update",
        "processes": "Complete Trip, Compute Round Trip Mileage, Update Truck Fuel",
        "triggered": "Driver, System",
        "fields": [
            ("truckTripUsageID", "Number", "PK", "Auto-number", "Unique truck trip usage identifier."),
            ("tripID", "Number", "Logical key", "N/A", "Related trip."),
            ("truckID", "Number", "FK/logical", "N/A", "Related truck."),
            ("oneWayDistanceKm", "Number", "", "0.00", "One-way trip distance."),
            ("roundTripDistanceKm", "Number", "", "0.00", "Round-trip distance."),
            ("efficiencyKmPerLiter", "Number", "", "0.00", "Fuel efficiency used for estimate."),
            ("fuelUsed", "Number", "", "0.00", "Fuel used by trip."),
            ("fuelBefore", "Number", "", "0.00", "Fuel before trip."),
            ("fuelAfter", "Number", "", "0.00", "Fuel after trip."),
            ("mileageBefore", "Number", "", "0.00", "Mileage before trip."),
            ("mileageAfter", "Number", "", "0.00", "Mileage after trip."),
            ("dateCreated", "Date/Time", "", "MM/DD/YYYY HH:MM", "Usage record creation date."),
        ],
    },
    "userrights": {
        "description": "The userrights entity is a legacy table used for older credential or rights records. It is currently documented as a standalone table because it has no enforced database relationship.",
        "source": "Legacy Admin/User Rights Record",
        "processes": "Legacy Login/Access Reference",
        "triggered": "Admin",
        "fields": [
            ("id", "Number", "PK", "Auto-number", "Unique record identifier."),
            ("userid", "Short Text(10)", "", "N/A", "Legacy user identifier."),
            ("empid", "Short Text(10)", "", "N/A", "Legacy employee identifier."),
            ("username", "Short Text(20)", "", "N/A", "Legacy username."),
            ("upassword", "Short Text(20)", "", "N/A", "Legacy password."),
        ],
    },
}


def add_title_pages(doc):
    add_centered(doc, "ALMODIEL TRUCKING SERVICES SYSTEM", bold=True, size=12)
    doc.add_paragraph()
    doc.add_paragraph()
    add_centered(doc, "An Information Management System Proposal")
    add_centered(doc, "Presented to the faculty of the")
    add_centered(doc, "Department of Computer Science")
    add_centered(doc, "NU Bacolod Incorporated")
    doc.add_paragraph()
    add_centered(doc, "In Partial Fulfillment")
    add_centered(doc, "of the Requirements for the Course")
    add_centered(doc, "Information Management (CTINFMGL)")
    doc.add_paragraph()
    add_centered(doc, "by")
    add_centered(doc, "Jethro T. Almodiel")
    add_centered(doc, "Arldrich A. Marcelino")
    add_centered(doc, "2026")
    doc.add_page_break()

    add_heading(doc, "LIST OF FIGURES", 1)
    figures = [
        "Figure 1. Entity-Relationship Diagram of Almodiel Trucking Services System",
        "Figure 2. Landing Page Interface",
        "Figure 3. Login Page Interface",
        "Figure 4. Admin Dashboard Interface",
        "Figure 5. Booking Registration Interface",
        "Figure 6. Trips Page Interface",
        "Figure 7. Trip Details Interface",
        "Figure 8. Customer Registration Interface",
        "Figure 9. Employee Registration Interface",
        "Figure 10. Truck Registration Interface",
        "Figure 11. Sales Page Interface",
        "Figure 12. Incident Reports Interface",
    ]
    for fig in figures:
        add_body(doc, fig)
    doc.add_page_break()

    add_heading(doc, "LIST OF TABLES", 1)
    table_names = [
        "Table 1. User Roles",
        "Table 2. Main Relationship Summary",
        "Table 3. Booking Entity",
        "Table 4. Cargo Entity",
        "Table 5. Customer Entity",
        "Table 6. Delivery Charge Entity",
        "Table 7. Employee Entity",
        "Table 8. Expenses Entity",
        "Table 9. Incident Report Entity",
        "Table 10. Location Entity",
        "Table 11. Sales Entity",
        "Table 12. Staff Salary Entity",
        "Table 13. Tariff Entity",
        "Table 14. Trip Employee Entity",
        "Table 15. Truck Entity",
        "Table 16. Truck Employee Entity",
        "Table 17. Truck Fuel Log Entity",
        "Table 18. Truck Trip Usage Entity",
        "Table 19. User Rights Entity",
        "Table 20. User Roles and Access Control",
    ]
    for table_name in table_names:
        add_body(doc, table_name)
    doc.add_page_break()

    add_heading(doc, "TABLE OF CONTENTS", 1)
    toc = [
        "CHAPTER 1 INTRODUCTION",
        "1.1 Organizational Background",
        "1.2 Problem Statements",
        "1.3 System Objectives",
        "1.4 Significance of the Study",
        "1.5 Definition of Terms",
        "CHAPTER 2 FORMS AND DATA ANALYSIS",
        "2.1 Form Description",
        "2.2 Entity Relationship Diagram",
        "2.3 Data Dictionary",
        "2.4 User Roles and Access Control",
        "CHAPTER 3 PROJECT TECHNICALITY",
        "3.1 System Project Scope",
        "3.2 Software Interface Descriptions",
    ]
    for item in toc:
        add_body(doc, item)
    doc.add_page_break()


def chapter_1(doc):
    add_heading(doc, "CHAPTER 1", 1)
    add_heading(doc, "INTRODUCTION", 1)
    intro = [
        "In today's organizations and institutions, effective information management is essential in supporting daily operations, decision-making, and quality service delivery. Accurate and organized records allow an organization to monitor transactions, track resources, and provide timely responses to the needs of its stakeholders. As the number of records and transactions increases, manual methods become more difficult to maintain because they require more time, physical storage, and human effort. For logistics and trucking operations, information management is especially important because booking records, delivery schedules, customer details, truck assignments, driver assignments, and billing information must be accurate and accessible.",
        "Almodiel Trucking Services currently handles trucking and delivery-related information that includes customer records, booking details, cargo descriptions, pickup and destination locations, delivery prices, truck information, driver and assistant assignments, delivery status, expenses, salary records, and billing reports. In a manual or partially manual setup, these records may be written in logbooks, encoded in separate files, or verified through physical documents and receipts. This process can support basic operations, but it becomes inefficient when several bookings, trucks, and delivery personnel need to be monitored at the same time. The staff, drivers, assistants, customers, and management all depend on reliable information to complete their respective responsibilities.",
        "The use of manual records can lead to delayed retrieval, inconsistent entries, duplicate information, misplaced documents, and difficulty in preparing reports. Delivery tracking can also become unclear when drivers and administrators do not share a centralized view of trip status. Billing, expenses, salary computation, and truck fuel or mileage monitoring may require repeated manual checking, which increases the risk of inaccurate figures. These issues affect the organization's operational efficiency, reporting accuracy, customer service, and ability to make timely decisions.",
        "To address these challenges, the proposed Almodiel Trucking Services System is designed as a web-based information management system that centralizes operational records and automates important transactions. The system supports customer registration, booking creation, multiple cargo entries, truck and crew assignment, map-based pickup and destination records, trip monitoring, delivery status updates, incident reporting, tariff management, sales monitoring, expenses, staff salary records, and truck fuel and mileage tracking. It aims to improve the accuracy, accessibility, and organization of data used in daily trucking operations. This proposal documents the background, problems, objectives, data requirements, user roles, project scope, and software interfaces of the current system."
    ]
    for p in intro:
        add_body(doc, p)

    add_heading(doc, "1.1 Organizational Background", 2)
    org = [
        "Almodiel Trucking Services is an organization engaged in providing trucking and delivery services to customers and partner companies. Its operation involves accepting bookings, assigning trucks and crew, transporting cargo, monitoring delivery progress, and preparing billing records for completed services. The organization depends on administrative staff, drivers, and assistants to coordinate deliveries and ensure that goods are transported properly from pickup point to destination.",
        "The organization manages several kinds of information, including customer details, company records, booking schedules, cargo information, delivery prices, truck records, driver and assistant assignments, location addresses, fuel records, mileage readings, expenses, salary records, and incident reports. These data are important because they serve as the basis for delivery planning, billing preparation, personnel monitoring, and management reporting. Without a centralized system, these records may be difficult to retrieve, verify, and update, especially when the number of transactions increases.",
        "The proposed system supports Almodiel Trucking Services by providing a structured digital platform for recording and monitoring these operational data. Instead of relying only on manual documents or separate records, the system allows authorized users to manage bookings, trips, reports, and records through a unified web-based interface. Proper data management is important for the organization because it improves accountability, reduces errors, supports faster service, and helps management make informed operational decisions."
    ]
    for p in org:
        add_body(doc, p)

    add_heading(doc, "1.2 Problem Statements", 2)
    problems = [
        ("1.2.1 Inefficient Manual Record Management", "The organization may experience difficulty in managing booking, customer, truck, and delivery records when information is handled through logbooks, paper forms, or disconnected files. This problem is caused by the absence of a centralized database where records can be stored, searched, updated, and retrieved quickly. As a result, staff may spend unnecessary time checking previous transactions, verifying delivery information, and locating important records. This reduces productivity and increases the possibility of incomplete or inconsistent information."),
        ("1.2.2 Difficulty in Monitoring Truck Availability and Trip Assignments", "Truck scheduling and crew assignment can become difficult when availability is checked manually. This problem occurs because administrators must compare booking dates, truck status, driver assignments, and assistant assignments before finalizing a trip. If these records are not centralized, there is a risk of assigning a truck or crew member to conflicting trips. This can cause delays, confusion, and inefficient use of company resources."),
        ("1.2.3 Slow Billing, Expense, and Salary Report Preparation", "Preparing billing statements, sales reports, expense summaries, and salary records can be time-consuming when values are computed or gathered manually. This problem is caused by scattered transaction records and the need to repeatedly verify booking prices, delivery charges, expenses, and staff salary data. The effect is slower report generation and a higher chance of inaccurate totals. This can affect management decisions and customer billing accuracy."),
        ("1.2.4 Limited Delivery Tracking and Incident Visibility", "Drivers and administrators need a reliable way to monitor trip status and report delivery problems. Without a digital status update and incident reporting feature, the organization may not immediately know whether a trip is pending, in transit, at stopover, completed, or affected by an incident. This problem can delay response to vehicle breakdowns, cargo concerns, route issues, or customer-related problems. The lack of timely visibility can affect customer service and operational control."),
        ("1.2.5 Inconsistent Truck Fuel, Mileage, and Maintenance Monitoring", "Truck fuel, mileage, and usage records are important in logistics operations but may be difficult to maintain manually. This problem is caused by separate recording of fuel entries, trip distances, maintenance expenses, and truck status updates. If these data are not connected to trips and truck records, it becomes harder to estimate fuel usage, monitor mileage, and prepare maintenance-related reports. This can lead to inaccurate operational costing and poor resource planning."),
    ]
    for heading, text in problems:
        add_heading(doc, heading, 3)
        add_body(doc, text)

    add_heading(doc, "1.3 System Objectives", 2)
    add_body(doc, "The proposed system aims to provide a centralized and organized information management platform for Almodiel Trucking Services. The objectives below directly address the identified operational problems and are translated into measurable system functions.")
    objectives = [
        ("1.3.1 Centralize Operational Record Management", "To develop a centralized database that stores customer, booking, cargo, location, truck, employee, trip, sales, expense, salary, and incident records. This objective addresses inefficient manual record management by allowing authorized users to add, update, search, and retrieve records through the system. The centralized structure reduces duplicate entries and improves data consistency. It also allows the organization to maintain a more reliable source of operational information."),
        ("1.3.2 Automate Booking, Truck Availability, and Crew Assignment", "To provide booking features that validate pickup date, truck availability, driver assignment, assistant assignment, cargo details, and location pins before saving a booking. This objective addresses scheduling and assignment problems by helping administrators avoid truck conflicts and incomplete crew assignments. The system supports trip grouping through `tripID` and connects bookings to assigned trucks and employees. This improves planning and reduces the risk of operational delays."),
        ("1.3.3 Generate Accurate Billing, Sales, Expense, and Salary Records", "To automate the organization of billing, sales, expenses, and staff salary data based on completed bookings and recorded operational costs. This objective addresses slow report preparation by allowing the system to compile data from connected tables instead of relying on manual computation. The system supports sales monitoring, payment status updates, report filtering, and salary records for drivers and assistants. This improves financial accuracy and report availability."),
        ("1.3.4 Provide Delivery Status Tracking and Incident Reporting", "To allow drivers to update assigned trips from pending to in-transit, stopover, and completed while also providing an incident report feature. This objective addresses limited delivery visibility by giving administrators and drivers a shared view of delivery progress. Incident reports can be reviewed by admins and resolved with notes or action taken. This supports faster response and better accountability."),
        ("1.3.5 Monitor Truck Fuel, Mileage, and Usage", "To provide truck detail pages, fuel logs, and trip usage records that track fuel and mileage before and after completed trips. This objective addresses inconsistent truck monitoring by connecting truck usage to delivery completion. The system can estimate round-trip mileage and fuel consumption for completed trips. This helps the organization monitor truck status and support maintenance planning."),
    ]
    for heading, text in objectives:
        add_heading(doc, heading, 3)
        add_body(doc, text)

    add_heading(doc, "1.4 Significance of the Study", 2)
    stakeholders = [
        ("1.4.1 Management and Administrators", "The system is significant to management and administrators because it provides an organized platform for monitoring bookings, trips, trucks, employees, billing, expenses, and incident reports. It reduces the time needed to retrieve records and prepare reports. Administrators can make better decisions because operational data are centralized and easier to verify. The system also improves accountability by recording who created or updated important records."),
        ("1.4.2 Drivers and Assistants", "The system benefits drivers and assistants by giving them clearer access to assigned trips and delivery details. Drivers can view pickup and destination information, update delivery status, and submit incident reports when problems occur. Salary records can also be viewed by drivers, making compensation information more transparent. This improves communication between field personnel and administrative staff."),
        ("1.4.3 Customers and Partner Companies", "Customers and partner companies benefit from more organized booking and delivery records. The system supports customer booking, booking details, and status visibility, which can improve trust and service quality. Accurate billing and route information also help reduce confusion during payment and delivery coordination. As a result, customers receive a more reliable service experience."),
        ("1.4.4 Future Developers and Researchers", "The documentation and system design can help future developers understand how trucking operations can be modeled through a web-based information management system. The database, ERD, and data dictionary provide references for improving or extending the project. Researchers may also use the study as a guide for similar logistics, transportation, or record management systems. This supports continuous improvement and future system development."),
    ]
    for heading, text in stakeholders:
        add_heading(doc, heading, 3)
        add_body(doc, text)

    add_heading(doc, "1.5 Definition of Terms", 2)
    terms = [
        ("Administrator", "A system user with full access to operational records. In the system, the administrator manages customers, employees, trucks, bookings, tariffs, reports, sales, and incident records."),
        ("Assistant", "A crew member assigned to support a driver during a trip. In the system, assistants are recorded in trip crew assignments and can view assigned trip information."),
        ("Booking", "A transaction record that represents a delivery request. In the system, it stores customer, pickup, destination, cargo, price, schedule, and status information."),
        ("Cargo", "The goods or items to be transported by the trucking service. In the system, cargo records are attached to bookings and may include multiple items."),
        ("Customer", "An individual or company requesting trucking services. In the system, customers can be linked to bookings, locations, tariffs, and sales records."),
        ("Delivery Charge", "An additional amount charged for hauling or other delivery-related costs. In the system, delivery charges are stored separately for reporting and billing support."),
        ("Driver", "An employee assigned to operate a truck and complete deliveries. In the system, drivers can view assigned trips, update delivery status, and submit incident reports."),
        ("Entity Relationship Diagram", "A database diagram that shows entities, fields, and relationships. In the system, the ERD documents how booking, customer, truck, employee, sales, and related records connect."),
        ("Incident Report", "A record of an accident, delay, breakdown, route issue, cargo concern, or other delivery problem. In the system, drivers submit incident reports and admins review them."),
        ("Location Pin", "A coordinate-based map marker representing a pickup or destination point. In the system, location pins are used to display routes and validate delivery addresses."),
        ("Sales Record", "A financial record generated from completed bookings. In the system, it stores gross amount, expenses, net amount, paid amount, balance, and payment status."),
        ("Tariff", "A pricing rule based on customer, route, truck type, distance, and fuel range. In the system, tariffs help calculate booking prices for company customers."),
        ("Trip", "A delivery assignment that may contain one or more bookings. In the system, trips are represented by the logical group value `booking.tripID`."),
        ("Truck Fuel Log", "A record of fuel and odometer information for a truck. In the system, it supports monitoring fuel usage and operational cost."),
        ("Truck Trip Usage", "A record of mileage and fuel usage after a completed trip. In the system, it supports automatic truck mileage and fuel updates."),
    ]
    for term, definition in terms:
        add_bullet(doc, term, definition)


def chapter_2(doc):
    doc.add_page_break()
    add_heading(doc, "CHAPTER 2", 1)
    add_heading(doc, "FORMS AND DATA ANALYSIS", 1)
    add_body(doc, "This chapter presents the forms and data analysis requirements of the Almodiel Trucking Services System. It identifies the main input forms used for data entry and record maintenance, explains the database structure through the entity relationship diagram, and documents all database fields through the data dictionary. It also discusses user roles and access control to show how system functions are organized according to responsibility.")

    add_heading(doc, "2.1 Form Description", 2)
    add_body(doc, "This section presents the major forms used in the proposed system. Each form is described according to its purpose, users, data captured, and the process it improves.")
    forms = [
        ("2.1.1 Customer Registration Form", "The Customer Registration Form is used to encode individual and company customer information. It captures customer type, name, contact details, account information, company documents, and location information. This form replaces scattered manual customer records by storing customer data in a structured database. After submission, the customer record becomes available for booking, tariff assignment, and customer-related reports."),
        ("2.1.2 Employee Registration Form", "The Employee Registration Form is used by the administrator to register drivers, assistants, and admin users. It records personal information, contact details, role, password, status, and license information for drivers. This form improves personnel management by keeping employee credentials and role classification in one table. After submission, employees can be assigned to trucks, trips, salary records, and incident reports."),
        ("2.1.3 Truck Registration Form", "The Truck Registration Form is used to create truck records and assign default crew members. It captures plate number, truck type, capacity, fuel, mileage, brand, documents, status, driver, and assistants. This form improves the manual process of monitoring truck information by linking trucks to assigned employees. After submission, the truck can be selected during booking and monitored in the truck management module."),
        ("2.1.4 Booking Registration Form", "The Booking Registration Form is the main transaction form for creating delivery bookings. It records the customer, store name, pickup date and time, truck, driver, assistants, crew salary, allowance, cargo items, hauling amount, price, pickup location, and destination location. It improves the manual booking process by validating required fields, truck availability, map coordinates, and cargo details before saving. After submission, the booking appears in the trips module and related records are created for cargo, crew, charges, and salary when applicable."),
        ("2.1.5 Trip Details and Delivery Status Form", "The Trip Details interface is used to view connected bookings, crew, route map, and delivery status actions. Drivers can start delivery, mark stopover, and mark delivered for assigned trips, while admins can modify trip information. This form improves delivery monitoring by replacing verbal or manual status updates with system-recorded progress. After status updates, the system updates booking status and can generate sales or truck usage records when the trip is completed."),
        ("2.1.6 Incident Report Form", "The Incident Report Form is used by drivers to report accidents, breakdowns, cargo damage, delays, route issues, customer issues, or other delivery concerns. It records trip, booking, incident type, severity, date and time, location, description, and action taken. This form improves incident visibility by providing a structured report that admins can review. After submission, the report appears in the admin incident reports page for review, resolution, dismissal, or reopening."),
        ("2.1.7 Tariff Management Form", "The Tariff Management Form is used by admins to define pricing rules for customers and company routes. It stores customer, branch, origin, destination, distance, truck type, base rate, fuel range, fuel subsidy, and status. This form improves price consistency by allowing the booking module to match rates based on route and fuel pump price. After submission, tariffs can be used during booking price calculation."),
        ("2.1.8 Truck Details and Fuel Log Form", "The Truck Details and Fuel Log interface is used to review truck status, mileage, fuel, assigned crew, fuel logs, and trip usage. It captures fuel entries, odometer readings, amount, station, reference number, and notes. This form improves truck monitoring by connecting fuel and mileage records to a specific truck. After submission, fuel log records are added to the truck history."),
        ("2.1.9 Sales and Billing Form", "The Sales and Billing interface is used by admins to review completed booking sales, payment status, paid amount, balance, trends, and date filters. It supports individual and grouped payment monitoring. This form improves billing control by organizing financial records generated from completed trips. After payment updates, the sales record reflects the current billing status."),
        ("2.1.10 Reports Form", "The Reports interface is used to review billing, expenses, staff records, and staff salary summaries. It provides filters for report category, specific report type, and date range. This form improves management reporting by allowing administrators to view operational and financial summaries in one area. After filters are applied, the system displays the relevant report records for review."),
    ]
    for heading, text in forms:
        add_heading(doc, heading, 3)
        add_body(doc, text)

    add_heading(doc, "2.2 Entity Relationship Diagram", 2)
    add_body(doc, "The Entity Relationship Diagram of the Almodiel Trucking Services System serves as the blueprint of the database structure. It identifies the major entities such as booking, customer, cargo, employee, truck, location, tariff, sales, expenses, staffsalary, incidentreport, tripemployee, truckemployee, truckfuellog, and trucktripusage. The ERD shows how records are connected to support booking creation, trip monitoring, billing, incident reporting, salary generation, and truck usage monitoring.")
    add_body(doc, "A key design characteristic of the system is that there is no separate physical trip table. Instead, the system groups one or more bookings through the `booking.tripID` field. Related records such as trip employees, truck trip usage, salary records, sales, expenses, delivery charges, and incident reports use this logical trip value to connect records to the same delivery assignment. This structure supports the current workflow while still allowing future improvement through a dedicated trip table.")
    add_body(doc, "Junction and associative tables are used to support relationships that involve multiple records. For example, `tripemployee` connects trips, trucks, and employees so that one trip can have a driver and assistants. The `truckemployee` table records default truck crew assignments, while `cargo` allows a booking to contain multiple cargo items. These relationships help reduce duplicate data, improve data consistency, and support reporting.")
    if ERD_IMAGE.exists():
        doc.add_page_break()
        section = doc.add_section()
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = None
        run = p.add_run()
        run.add_picture(str(ERD_IMAGE), width=Inches(10.4))
        add_caption(doc, "Figure 1. Entity-Relationship Diagram of Almodiel Trucking Services System")
        doc.add_page_break()
        section = doc.add_section()
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    add_heading(doc, "2.3 Data Dictionary", 2)
    add_body(doc, "The data dictionary presents the structure of each entity included in the ERD. Each table identifies the source document or form, entity description, processes where it is used, triggering user, fields, data types, keys, formats, and field descriptions.")
    table_no = 3
    widths = [2200, 2300, 1200, 2600, 4100]
    for table_name, meta in TABLES.items():
        add_heading(doc, f"2.3.{table_no - 2} {table_name} Entity", 3)
        add_body(doc, f"The `{table_name}` entity is documented in Table {table_no}. {meta['description']} Its source document is {meta['source']}. The process or processes where this entity is used include {meta['processes']}, and the transaction is triggered by {meta['triggered']}.")
        add_caption(doc, f"Table {table_no}. {table_name} Entity")
        add_table(
            doc,
            ["Field Name", "Type", "Key", "Format", "Description"],
            field_rows(meta["fields"]),
            widths=widths,
        )
        table_no += 1

    add_heading(doc, "2.4 User Roles and Access Control", 2)
    add_body(doc, "This section discusses how access to the system is organized according to user roles. Role-based access control helps ensure that users only perform actions related to their responsibilities and prevents unauthorized changes to critical records.")
    roles = [
        ("2.4.1 Administrator", "The Administrator is the main system user responsible for managing operational and financial records. This user can access the dashboard, registrations, booking form, trips, trip details, management records, tariff records, truck details, sales, reports, and incident reports. The administrator can register employees, customers, and trucks, create and modify bookings, review sales, manage tariffs, and resolve incident reports. The administrator should not use the driver status flow unless acting in an authorized management role because delivery progress is primarily a driver responsibility. This role supports proper system control because administrative actions are limited to authorized personnel."),
        ("2.4.2 Driver", "The Driver is the delivery user assigned to transport cargo for specific trips. This user can view assigned trips, open trip details, start delivery, mark stopover, mark delivered, submit incident reports, and view salary records. The driver cannot manage customer records, employee records, tariffs, sales, or system-wide reports. This role is important because it separates field delivery actions from administrative management actions. Access control ensures that a driver only updates trips assigned to the driver's account."),
        ("2.4.3 Assistant", "The Assistant is a crew member assigned to support a driver during a delivery. This user can view assigned trip information through the assistant-accessible trip pages. The assistant cannot create bookings, update billing, manage trucks, resolve incidents, or modify administrative records. This limitation prevents unauthorized changes while still allowing assistants to view information needed for delivery support. The role improves operational coordination without giving unnecessary access to sensitive records."),
        ("2.4.4 Customer Individual", "The Customer Individual role represents a personal customer account. This user can register, create bookings, view personal bookings, and open booking details. The customer cannot manage other customers, employees, trucks, tariffs, reports, or sales records. This role supports self-service booking while protecting internal company data. Access control ensures that customers only see records related to their own account."),
        ("2.4.5 Customer Company", "The Customer Company role represents company customers or partner companies that request trucking services. This user can access customer booking functions and view company-related booking details. The company customer cannot modify employee, truck, tariff, sales, or incident records unless handled by an administrator. This role supports organized booking transactions for company clients. It also allows the system to connect company bookings with tariff-based pricing and warehouse pickup information."),
    ]
    for heading, text in roles:
        add_heading(doc, heading, 3)
        add_body(doc, text)
    add_caption(doc, "Table 20. User Roles and Access Control")
    add_table(doc, ["User Role", "Allowed Access", "Restricted Access"], [
        ["Administrator", "Dashboard, records management, booking, trips, sales, reports, incident review, tariff and truck management", "None within normal system administration, but must follow organizational authorization"],
        ["Driver", "Assigned trips, trip details, delivery status updates, incident submission, salary view", "Admin records, sales management, tariff management, customer/employee/truck management"],
        ["Assistant", "Assigned trip views and assistant dashboard", "Delivery status update, admin records, reports, sales, incident resolution"],
        ["Customer Individual", "Own booking form, own bookings, booking details, profile", "Internal operations, other customer records, sales reports, truck and employee records"],
        ["Customer Company", "Company booking access and company-related booking details", "Internal operations, employee/truck management, incident review, sales administration"],
    ], widths=[2600, 5600, 5600])


def chapter_3(doc):
    doc.add_page_break()
    add_heading(doc, "CHAPTER 3", 1)
    add_heading(doc, "PROJECT TECHNICALITY", 1)
    add_heading(doc, "3.1 System Project Scope", 2)
    add_body(doc, "The Almodiel Trucking Services System covers the major information management processes needed for trucking operations. It includes user authentication, customer and company records, booking creation, cargo recording, truck and crew assignment, trip monitoring, delivery status updates, incident reporting, tariff management, billing, sales, expenses, salary records, reports, and truck fuel and mileage monitoring. The system is intended for web-based use by authorized administrators, drivers, assistants, and customers.")
    modules = [
        ("3.1.1 Authentication and Role-Based Access Module", "This module manages system login and user access according to role. It allows admins, drivers, assistants, and customers to access only the pages and actions assigned to them. The module uses session data to determine whether a user is logged in and what routes are available. It prevents unauthorized users from opening restricted pages. This module solves access control problems by organizing system privileges based on user responsibility."),
        ("3.1.2 Customer and Company Management Module", "This module manages individual and company customer records. Administrators can register, view, and maintain customer information, while customers can use account-related pages to access their own records. The module stores contact details, customer type, company documents, and location information. It supports booking creation because every booking must be connected to a customer. This module improves record organization and customer data retrieval."),
        ("3.1.3 Booking and Cargo Management Module", "This module allows users to create booking records with customer, schedule, truck, crew, cargo, price, hauling, pickup, and destination data. It supports multiple cargo entries for one booking and validates required inputs before saving. The module checks truck availability and blocks past pickup dates. It also stores pickup and destination location records with coordinates for mapping. This module addresses manual booking problems by automating the main delivery request process."),
        ("3.1.4 Trip and Delivery Monitoring Module", "This module displays trips, connected bookings, route maps, assigned crew, and delivery status. Drivers can start delivery, mark stopover, and mark delivered for assigned trips. Admins can view all trips and modify selected trip details when needed. The module updates booking status and supports sales and truck usage updates after completion. This module improves delivery visibility and operational monitoring."),
        ("3.1.5 Truck Management and Fuel Monitoring Module", "This module manages truck registration, truck status, capacity, fuel, mileage, default crew, documents, fuel logs, and trip usage history. Administrators can view truck details and monitor fuel and mileage changes. The system can update truck fuel and mileage after completed trips using round-trip distance estimates. This module supports maintenance planning and operational costing. It reduces inconsistent truck monitoring by centralizing truck-related records."),
        ("3.1.6 Employee and Crew Assignment Module", "This module manages employee registration and crew assignment. It stores admins, drivers, and assistants, including contact details, role, status, and license information. During booking, the system assigns a driver and assistants to a trip through the trip employee record. This module supports driver trip filtering and salary generation. It improves personnel coordination and accountability."),
        ("3.1.7 Tariff and Pricing Module", "This module stores route-based pricing records for customers and company accounts. Admins can define origin, destination, truck type, distance, base rate, fuel ranges, and fuel subsidy. The booking module can use tariff data to support price calculation when the route and fuel pump price match an active tariff. This module improves pricing consistency and reduces manual computation. It is especially useful for repeat company customers."),
        ("3.1.8 Sales, Billing, and Payment Module", "This module organizes billing and sales records from completed bookings. Admins can view sales summaries, payment status, paid amount, balance, and date-based filters. The module supports grouped payment monitoring for periodic billing. It also deducts expenses to show net sales values. This module addresses slow billing and report preparation by using system-generated financial data."),
        ("3.1.9 Expense, Report, and Salary Module", "This module manages operational expenses, staff salary records, billing reports, staff lists, and salary reports. Admins can review expense categories such as fuel, maintenance, salary, tolls, parking, repair, office costs, and other expenses. Salary records can be generated from bookings and viewed by drivers. Reports provide organized outputs for management review. This module improves financial monitoring and payroll transparency."),
        ("3.1.10 Incident Report and Notification Module", "This module allows drivers to submit incident reports and allows admins to review and resolve them. Incident records include trip, booking, incident type, severity, date/time, location, description, action taken, and admin notes. The admin notification area can show incident-related updates along with upcoming and completed trip notifications. This module improves safety monitoring and response accountability. It gives management a clearer view of delivery problems."),
    ]
    for heading, text in modules:
        add_heading(doc, heading, 3)
        add_body(doc, text)

    add_heading(doc, "3.2 Software Interface Descriptions", 2)
    add_body(doc, "This section describes the major software interfaces of the Almodiel Trucking Services System. The interfaces are presented in a logical order based on user access and operational workflow. Each interface supports a specific part of the system, such as registration, booking, monitoring, reporting, or management.")
    interfaces = [
        ("3.2.1 Landing Page Interface", "The Landing Page Interface introduces Almodiel Trucking Services to public users. It presents service information, company details, and navigation to login or registration options. Users can read about the service and proceed to the appropriate login or signup page. This interface supports first-time access and public information browsing."),
        ("3.2.2 Login Page Interface", "The Login Page Interface allows admins, drivers, assistants, and customers to access the system using their credentials. It validates the phone number or account information and password before creating a session. When login is successful, the system redirects the user based on role. If the credentials are incorrect, the system displays an error message."),
        ("3.2.3 Admin Dashboard Interface", "The Admin Dashboard Interface provides a summary of key operational information such as sales, upcoming trips, truck status, and schedule. It includes dashboard cards and calendar-style schedule indicators. Admins use this page to quickly understand current operations. Notification badges help alert the admin to upcoming trips, completed deliveries, and incident reports."),
        ("3.2.4 Booking Registration Interface", "The Booking Registration Interface is used to create delivery bookings. It includes customer selection, store/customer name, pickup date and time, truck, driver, assistants, crew salary, allowance, cargo fields, hauling amount, price, and map-based pickup and destination inputs. The user can add multiple cargo items and validate route information. Clicking the save button records the booking and related information if all requirements are satisfied."),
        ("3.2.5 Trips Page Interface", "The Trips Page Interface lists trips based on the current user's role. Admins can view all trips, while drivers and assistants view assigned trips. The interface shows schedule, booking information, pickup, destination, status, and action buttons. Selecting a trip opens or prepares the trip details view."),
        ("3.2.6 Trip Details Interface", "The Trip Details Interface shows connected bookings, customer information, route map, crew, trip distance, and delivery status. Drivers can start delivery, mark stopover, and mark delivered when allowed by the current status. Admins can modify trip schedule, truck, crew, selected booking destination, and booking price. The map displays pickup and destination pins for route visibility."),
        ("3.2.7 Customer Registration Interface", "The Customer Registration Interface allows customer or company records to be added to the system. It captures customer type, name, contact person, phone number, email, documents, and location data. Admins use this interface to maintain company records, while customers may use signup pages for account creation. After saving, the customer can be selected during booking."),
        ("3.2.8 Employee Registration Interface", "The Employee Registration Interface allows administrators to register drivers, assistants, and admins. It includes fields for personal details, role, contact information, password, status, license number, license expiration, and license image. The form validates required information before saving. Registered employees can then be assigned to trucks and trips."),
        ("3.2.9 Truck Registration Interface", "The Truck Registration Interface is used to add truck records and assign default crew members. It includes truck plate number, type, capacity, fuel, mileage, brand, documents, status, driver, and assistant selections. Admins use this interface to maintain the fleet. After saving, the truck becomes available for booking assignment if active and not conflicting."),
        ("3.2.10 Manage Company Interface", "The Manage Company Interface allows admins to view and maintain company/customer records. It supports record review and update actions depending on the available management controls. This interface helps administrators keep customer data current. It also supports operational accuracy because bookings depend on correct customer information."),
        ("3.2.11 Manage Tariff Interface", "The Manage Tariff Interface allows admins to create and update tariff records. It includes customer, branch, origin, destination, distance, truck type, base rate, fuel range, fuel subsidy, and status. These records are used by the booking form to match pricing. The interface improves price consistency for company routes."),
        ("3.2.12 Manage Truck and Truck Details Interface", "The Manage Truck Interface lists registered trucks, while the Truck Details Interface shows a specific truck's current status, mileage, fuel, crew, fuel logs, and trip usage. Admins can use these interfaces to monitor fleet condition and usage history. Fuel log entries support operational costing. Truck usage records support mileage and fuel monitoring after completed trips."),
        ("3.2.13 Sales Interface", "The Sales Interface presents completed booking sales, billing status, payment status, paid amount, balance, and financial trends. Admins can filter records by date range and status. The interface supports individual and grouped payment monitoring. It helps management review revenue and outstanding balances."),
        ("3.2.14 Reports Interface", "The Reports Interface provides access to billing, expenses, staff list, and staff salary records. Admins can select report category, specific report, and date range. The interface displays organized records for management review. This page supports decision-making by gathering important operational and financial data in one place."),
        ("3.2.15 Incident Reports Interface", "The Incident Reports Interface allows admins to review reports submitted by drivers. It displays incident type, severity, trip, booking, driver, description, action taken, status, and admin notes. Admins can review, resolve, dismiss, or reopen reports based on the situation. This interface supports safety monitoring and incident accountability."),
        ("3.2.16 Driver Salary Interface", "The Driver Salary Interface allows drivers to view salary records connected to their trips. It shows route or booking details, pay period, role, gross pay, deductions, net pay, and status. This interface gives drivers visibility over their compensation records. It supports transparency between the organization and delivery personnel."),
        ("3.2.17 Customer Bookings and Booking Details Interface", "The Customer Bookings Interface allows customers to view their own booking records, while the Booking Details Interface shows detailed information for a selected booking. Customers can see booking status, pickup and destination details, map information, and cargo details. This interface improves customer access to delivery information. It also reduces the need for manual follow-up regarding booking progress."),
    ]
    for heading, text in interfaces:
        add_heading(doc, heading, 3)
        add_body(doc, text)


def main():
    DOCS.mkdir(exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_title_pages(doc)
    chapter_1(doc)
    chapter_2(doc)
    chapter_3(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
